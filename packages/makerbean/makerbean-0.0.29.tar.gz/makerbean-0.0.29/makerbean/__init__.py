# -*- coding: utf-8 -*-
# @Author: Anderson
# @Date:   2019-11-11 17:42:18
# @Last Modified by:   ander
# @Last Modified time: 2020-11-21 19:18:19
from openpyxl import Workbook, load_workbook
from docx import Document
import requests
from bs4 import BeautifulSoup
from urllib.parse import quote_plus
import pdfplumber
from PyPDF2 import PdfFileReader, PdfFileWriter, PdfFileMerger
import jieba
import os
import re
import json
import csv
from pyecharts import options as opts
from pyecharts.charts import Bar
from pyecharts.charts import WordCloud
from pyecharts.charts import Map3D
from pyecharts.globals import ChartType
from pyecharts.commons.utils import JsCode
from copy import copy
from collections import Counter
from .stop_words import stop_words
from .data import provinces_coordinates


def validate_title(title):
	rstr = r"[\/\\\:\*\?\"\<\>\|\%]"
	new_title = re.sub(rstr, "_", title)
	return new_title


def mkdir(folder):
	if not os.path.exists(folder):
		os.mkdir(folder)


class WebCrawlerBot(object):
	"""docstring for WebCrawlerBot"""

	def __init__(self):
		self.headers = {
			'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/72.0.3626.119 Safari/537.36',
		}
		self.weibo_hot_list = []
		self.liepin_urls = {}
		self.抓取猎聘 = self.get_liepin
		self.抓取论文 = self.get_arxiv

	def set_cookie(self, cookie):
		self.headers['cookie'] = cookie

	def get_sum_box_office(self):
		req = requests.post("http://cbooo.cn/API/GetData.ashx", data={'MethodName': 'BoxOffice_GetPcHomeList'})
		return req.json()["Data"]['Table'][0]['dapan']

	def get_rank_movie_data(self, index):
		req = requests.post("http://cbooo.cn/API/GetData.ashx", data={'MethodName': 'BoxOffice_GetPcHomeList'})
		if index < len(req.json()["Data"]['Table1']):
			movie_data = req.json()["Data"]['Table1'][index]
			return [
				movie_data["MovieName"],
				movie_data["boxoffice"],
				movie_data["amount"],
				movie_data["releasedate"],
				movie_data["default_url"],
			]
		else:
			raise Exception("排名数字超出范围")

	def get_weibo_hot(self, index):
		# 列表第一位不是正常内容
		index += 1
		if not self.weibo_hot_list:
			base_url = 'https://s.weibo.com/top/summary'
			req = requests.get(base_url, headers=self.headers)
			soup = BeautifulSoup(req.text, 'lxml')
			today_hot = soup.select('#pl_top_realtimehot tr')[1:]
			self.weibo_hot_list = copy(today_hot)
		item = self.weibo_hot_list[index]
		title = item.select('.td-02 a')[0].text.strip()
		hot_count = int(item.select('.td-02 span')[0].text.strip())
		url = item.select('.td-02 a')[0].get('href')
		if 'javascript' in url:
			url = item.select('.td-02 a')[0].get('href_to')
		url = f'https://s.weibo.com{url}'

		# Get detail info
		req = requests.get(url, headers=self.headers)
		soup = BeautifulSoup(req.text, "lxml")
		author = soup.select(".card-wrap .content .info .name")[0].text.strip()
		content = soup.select(".card-wrap .content .txt")[0].text.strip()
		return [hot_count, title, author, content, url]

	def get_liepin(self, keyword, start_page, end_page=None):
		if keyword not in self.liepin_urls:
			req = requests.get(
				f'https://www.liepin.com/zhaopin/?sfrom=click-pc_homepage-centre_searchbox-search_new&d_sfrom=search_fp&key={quote_plus(keyword)}',
				headers=self.headers)
			soup = BeautifulSoup(req.text, 'lxml')
			self.liepin_urls[keyword] = 'https://www.liepin.com'
			self.liepin_urls[keyword] += '&d_curPage={d_curPage}'.join(soup.select('.pagerbar a')[3]['href'].split('&d_curPage='))
			self.liepin_urls[keyword] = self.liepin_urls[keyword].split('&curPage=')[0] + '&curPage={curPage}'
		if end_page is None:
			end_page = start_page + 1
		results = []
		for page in range(start_page, end_page):
			if page > 0:
				url = self.liepin_urls[keyword].format(d_curPage=page - 1, curPage=page)
			else:
				url = self.liepin_urls[keyword].format(d_curPage=page + 1, curPage=page)
			req = requests.get(url, headers=self.headers)
			soup = BeautifulSoup(req.text, 'lxml')
			for item in soup.select('.sojob-item-main'):
				job_name = item.select('h3 a')[0].text.strip()
				job_company = item.select('.company-name')[0].text.strip()
				job_field = item.select('.field-financing')[0].text.strip()
				job_salary = item.select('.condition .text-warning')[0].text.strip()
				if job_salary == '面议':
					annual_salary = -1
				else:
					if '-' in job_salary:
						min_salary = int(job_salary[:job_salary.index('-')])
						max_salary = int(job_salary[job_salary.index('-') + 1:job_salary.index('k')])
						months = int(job_salary[job_salary.index('·') + 1:-1])
						annual_salary = (min_salary + max_salary) / 2 * months * 1000
					else:
						monthly_salary = int(job_salary.split('k')[0])
						months = int(job_salary[job_salary.index('·') + 1:-1])
						annual_salary = monthly_salary * months * 1000
				job_area = item.select('.condition .area')[0].text.strip()
				job_edu = item.select('.condition .edu')[0].text.strip()
				job_experience = item.select('.condition span')[-1].text.strip()
				results.append([job_name, job_company, job_field, job_salary, annual_salary, job_area, job_edu, job_experience])
		return results

	def get_huaban(self, keyword, page, key='k4rwsxf5'):
		url = f"https://huaban.com/search/?q={quote_plus(keyword)}&type=pins&{key}&page={page+1}&per_page=20&wfl=1"
		req = requests.get(url, headers=self.headers)
		source = str(req.text)
		start_index = source.index('app.page["pins"] = ') + len('app.page["pins"] = ')
		end_index = start_index + source[start_index:].index('app.page["page"]')
		results = []
		for img in json.loads(source[start_index:end_index].strip()[:-1]):
			results.append({
				'url': f"https://hbimg.huabanimg.com/{img['file']['key']}",
				'name': validate_title(f"{img['board']['title']}-{img['pin_id']}.jpg")
			})

		return results

	def get_arxiv(self, keyword, page):
		url = f"https://arxiv.org/search/?query={quote_plus(keyword)}&searchtype=all&source=header&order=-announced_date_first&size=50&abstracts=show&start={50*page}"
		req = requests.get(url)
		soup = BeautifulSoup(req.text, 'lxml')
		results = soup.select('.arxiv-result')
		output_results = []
		for result in results:
			output_result = {}
			output_result['title'] = result.select('.title')[0].text.strip()
			authors = []
			for author in result.select('.authors a'):
				authors.append(author.text.strip())
			output_result['authors'] = authors
			output_result['abstract'] = result.select('.abstract-full')[0].text.replace('△ Less', '').strip()
			if result.find('a', string='pdf'):
				output_result['pdf'] = result.find('a', string='pdf').attrs['href']
			else:
				output_result['pdf'] = ''

			output_results.append(output_result)

		return output_results.copy()

	def download_image(self, url, filename, folder):
		mkdir(folder)
		req = requests.get(url)
		with open(os.path.join(folder, f'{validate_title(filename)}'), 'wb') as f:
			f.write(req.content)


class ExcelBot(object):
	"""docstring for ExcelBot"""

	def __init__(self):
		self.workbook = Workbook()
		self.sheet = self.workbook.active
		self.添加数据 = self.add_data
		self.添加一行数据 = self.add_row
		self.提取一行数据 = self.get_row
		self.提取一列数据 = self.get_col
		self.保存文件 = self.save

	def add_data(self, data):
		try:
			for row in data:
				self.add_row(row)
		except Exception as e:
			print(e)

	def add_row(self, row):
		try:
			self.sheet.append(row)
		except Exception as e:
			print(e)

	def get_row(self, row):
		data = []
		if isinstance(row, int):
			for cell in self.sheet[row + 1]:
				data.append(cell.value)

		return data

	def get_col(self, col):
		data = []
		if isinstance(col, str):
			for cell in self.sheet[col]:
				data.append(cell.value)
		elif isinstance(col, int):
			# excel column counts from 1
			col += 1
			for row in self.sheet.iter_rows(min_col=col, max_col=col):
				data.append(row[0].value)

		return data

	def clear(self):
		self.workbook.remove_sheet(self.sheet)
		self.sheet = self.workbook.create_sheet('sheet1')

	def open(self, filename):
		if filename.endswith('.csv'):
			self.workbook = Workbook()
			self.sheet = self.workbook.active
			with open(filename, encoding='utf-8') as f:
				reader = csv.reader(f)
				for row in reader:
					self.sheet.append(row)
		else:
			self.workbook = load_workbook(filename=filename)
			self.sheet = self.workbook.active

	def save(self, filename, ext='xlsx'):
		self.filename = filename
		if ext == 'xlsx':
			self.workbook.save(filename=f'{filename}.xlsx')
		elif ext == 'csv':
			with open(f'{filename}.csv', 'w', newline="", encoding='utf-8') as f:
				c = csv.writer(f)
				for r in self.sheet.rows:
					c.writerow([cell.value for cell in r])


class WordBot(object):
	"""docstring for WordBot"""

	def __init__(self):
		self.doc = Document()

	def set_paragraph(self, index, text):
		if index < len(self.doc.paragraphs):
			self.doc.paragraphs[index].text = str(text)
			return True
		else:
			return False

	def get_paragraph(self, index):
		if index < len(self.doc.paragraphs):
			return str(self.doc.paragraphs[index].text)
		else:
			return ''

	@property
	def paragraphs(self):
		return [str(p.text) for p in self.doc.paragraphs]

	def clear(self):
		self.doc = Document()

	def open(self, filename):
		self.doc = Document(filename)

	def save(self, filename):
		if not filename:
			filename = 'tmp'
		self.doc.save(f'{filename}.docx')


class PDFBot(object):
	"""docstring for ExcelBot"""

	def __init__(self):
		self.page_num = 0

	def open(self, file_path):
		self.filename, _ = os.path.splitext(os.path.basename(file_path))
		self.pdf = pdfplumber.open(file_path)
		self.pdf_reader = PdfFileReader(file_path)
		self.page_num = self.pdf_reader.getNumPages()

	def get_text(self, page):
		pdf_page = self.pdf.pages[page]
		return pdf_page.extract_text()

	def split(self, page, folder):
		mkdir(folder)
		pdf_writer = PdfFileWriter()
		pdf_writer.addPage(self.pdf_reader.getPage(page))
		with open(os.path.join(folder, f'{self.filename}-p{page}.pdf'), 'wb') as out:
			pdf_writer.write(out)

	def merge(self, pdfs, merged_name):
		merger = PdfFileMerger()
		for pdf in pdfs:
			merger.append(PdfFileReader(pdf))
		merger.write(f"{merged_name}.pdf")


class DataAnalysisBot(object):
	"""docstring for DataAnalysisBot"""
	def __init__(self):
		self.data = []
		self.分析词语频次 = self.get_word_frequency
		self.生成词云图 = self.generate_word_cloud

	def set_data(self, data):
		self.data = copy(data)

	def get_word_frequency(self, data, count=20):
		word_frequency = []
		words = ''
		if isinstance(data, str):
			words = data
		elif isinstance(data, list):
			words = '\n'.join([str(item) for item in data])

		punct = set(u''' #:!),.:;?]}¢'"、。〉》」』】〕〗〞︰︱︳﹐､﹒
		﹔﹕﹖﹗﹚﹜﹞！），．：；？｜｝︴︶︸︺︼︾﹀﹂﹄﹏､～￠
		々‖•·ˇˉ―--′’”([{£¥'"‵〈《「『【〔〖（［｛￡￥〝︵︷︹︻
		︽︿﹁﹃﹙﹛﹝（｛“‘-—_…@~/\\''')

		words_cut = list(filter(lambda x: x not in punct, jieba.lcut(words)))
		words_cut = list(filter(lambda x: x not in stop_words, words_cut))
		word_frequency = Counter(words_cut).most_common(count)

		return copy(word_frequency)

	def generate_word_cloud(self, data):
		wordcloud = WordCloud()
		wordcloud.add("", data, word_size_range=[20, 100])
		wordcloud.render('word_cloud.html')

	def generate_bar(self, x_axis, y_axis):
		bar = Bar(
			init_opts=opts.InitOpts(
				width="1280px",
				height="720px"))
		bar.add_xaxis(x_axis)
		bar.add_yaxis("", y_axis)
		bar.set_global_opts(
			xaxis_opts=opts.AxisOpts(axislabel_opts=opts.LabelOpts(rotate=10)))
		bar.render('bar.html')

	def generate_3d_map(self, provinces_data):
		map_3d = Map3D(
			init_opts=opts.InitOpts(
				width="1280px",
				height="720px"))
		map_3d.add_schema(
			itemstyle_opts=opts.ItemStyleOpts(
				color="rgb(242, 105, 92)",
				opacity=1,
				border_width=0.8,
				border_color="rgb(242, 181, 145)",
			),
			map3d_label=opts.Map3DLabelOpts(
				is_show=False,
				formatter=JsCode("function(data){return data.name + " " + data.value[2];}"),
			),
			emphasis_label_opts=opts.LabelOpts(
				is_show=False,
				color="#fff",
				font_size=10,
				background_color="rgba(242, 149, 128, 0)",
			),
			light_opts=opts.Map3DLightOpts(
				main_color="#fff",
				main_intensity=1.2,
				main_shadow_quality="high",
				is_main_shadow=True,
				main_beta=10,
				ambient_intensity=0.3,
			),
			is_show_ground=True,
			ground_color="#F2B591",
		)

		map_3d_data = []
		for item in provinces_data.items():
			province_data = (
				item[0],
				[*provinces_coordinates[item[0]], item[1]]
			)
			map_3d_data.append(province_data)

		map_3d.add(
			series_name="bar3D",
			data_pair=map_3d_data,
			type_=ChartType.BAR3D,
			bar_size=1,
			itemstyle_opts=opts.ItemStyleOpts(
				color="rgb(173, 212, 217)"
			),
			shading="realistic",
			label_opts=opts.LabelOpts(
				is_show=False,
				formatter=JsCode("function(data){return data.name + ' ' + data.value[2];}"),
			),
		)
		map_3d.set_global_opts(title_opts=opts.TitleOpts(title="Map3D-Bar3D"))
		map_3d.render("map3d_with_bar3d.html")


爬虫机器人 = pachong = wc_bot = web_crawler_bot = WebCrawlerBot()
表格机器人 = biaoge = ec_bot = excel_bot = ExcelBot()
数据机器人 = shuju = da_bot = data_analysis_bot = DataAnalysisBot()
PDF机器人 = pdf_bot = PDFBot()
文档机器人 = word_bot = WordBot()
